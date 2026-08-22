import os
import fitz
from docx import Document as DocxDocument

from models import Document
from cleaner import TextCleaner


class DocumentLoader:

    def load(self, file_path: str):

        extension = os.path.splitext(file_path)[1].lower()

        if extension == ".pdf":
            return self.load_pdf(file_path)

        if extension == ".docx":
            return self.load_docx(file_path)

        if extension == ".txt":
            return self.load_txt(file_path)

        raise ValueError("Unsupported file type.")

    def create_metadata(
        self,
        source,
        doc_type,
        content,
        page=None
    ):

        return {
            "source": os.path.basename(source),
            "page": page,
            "doc_type": doc_type,
            "word_count": len(content.split()),
            "char_count": len(content)
        }
    def load_pdf(
        self,
        file_path
    ):

        documents = []

        pdf = fitz.open(file_path)

        for page_number, page in enumerate(pdf, start=1):

            text = page.get_text()

            text = TextCleaner.clean(text)

            metadata = self.create_metadata(
                source=file_path,
                doc_type="pdf",
                content=text,
                page=page_number
            )

            documents.append(
                Document(
                    content=text,
                    metadata=metadata
                )
            )

        pdf.close()

        return documents
        def load_docx(
        self,
        file_path
    ):

        doc = DocxDocument(file_path)

        text = "\n".join(
            paragraph.text
            for paragraph in doc.paragraphs
            if paragraph.text.strip()
        )

        text = TextCleaner.clean(text)

        metadata = self.create_metadata(
            source=file_path,
            doc_type="docx",
            content=text
        )

        return [
            Document(
                content=text,
                metadata=metadata
            )
        ]
    def load_txt(
        self,
        file_path
    ):

        with open(
            file_path,
            "r",
            encoding="utf-8"
        ) as file:

            text = file.read()

        text = TextCleaner.clean(text)

        metadata = self.create_metadata(
            source=file_path,
            doc_type="txt",
            content=text
        )

        return [
            Document(
                content=text,
                metadata=metadata
            )
        ]
            